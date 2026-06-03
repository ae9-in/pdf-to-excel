import pdfplumber
from pypdf import PdfReader
from pdf2image import convert_from_path
import pytesseract
from PIL import Image
import os
from pathlib import Path

# --- Windows Path Auto-Resolution ---
TESSERACT_WINDOWS_PATHS = [
    r"C:\Program Files\Tesseract-OCR\tesseract.exe",
    r"C:\Program Files (x86)\Tesseract-OCR\tesseract.exe",
    r"C:\Users\jishn\AppData\Local\Programs\Tesseract-OCR\tesseract.exe",
    r"D:\Tesseract-OCR\tesseract.exe"
]

# Configure pytesseract path if found in standard Windows locations
for p in TESSERACT_WINDOWS_PATHS:
    if os.path.exists(p):
        pytesseract.pytesseract.tesseract_cmd = p
        break

POPPLER_WINDOWS_PATHS = [
    r"C:\Program Files\poppler\bin",
    r"C:\Program Files\Poppler\bin",
    r"C:\Program Files\Poppler\Library\bin",
    r"C:\poppler\bin",
    r"D:\poppler\bin",
    r"C:\Users\jishn\AppData\Local\Programs\poppler\bin"
]

def get_poppler_path() -> str:
    """Returns the path to poppler bin folder if found, otherwise None."""
    # Check standard paths
    for p in POPPLER_WINDOWS_PATHS:
        if os.path.exists(p):
            return p
    # Also check inside directories in C:\Program Files starting with poppler
    try:
        pf = Path("C:/Program Files")
        if pf.exists():
            for sub in pf.iterdir():
                if sub.is_dir() and "poppler" in sub.name.lower():
                    bin_path = sub / "bin"
                    if bin_path.exists():
                        return str(bin_path)
                    lib_bin = sub / "Library" / "bin"
                    if lib_bin.exists():
                        return str(lib_bin)
    except Exception:
        pass
    return None


def detect_pdf_type(pdf_path: str) -> str:
    """
    Returns 'text' if the PDF has selectable text,
    'scanned' if it's an image-based PDF needing OCR.
    """
    try:
        with pdfplumber.open(pdf_path) as pdf:
            total_chars = 0
            # Check up to first 3 pages
            pages_to_check = pdf.pages[:3]
            for page in pages_to_check:
                text = page.extract_text() or ""
                total_chars += len(text.strip())
            
            avg_chars = total_chars / len(pages_to_check) if pages_to_check else 0
            # If less than 50 chars per page on average -> likely scanned
            return "text" if avg_chars > 50 else "scanned"
    except Exception as e:
        # Fallback to scanned if open fails
        return "scanned"


def extract_text_from_pdf(pdf_path: str) -> tuple[str, int]:
    """
    Extract text from a digitally-created PDF using pdfplumber.
    Uses layout-preserving text extraction which already captures tabular
    content in correct column order. We do NOT add table_text on top of that
    to avoid duplicating rows.
    """
    full_text = []
    page_count = 0

    with pdfplumber.open(pdf_path) as pdf:
        page_count = len(pdf.pages)
        for page_num, page in enumerate(pdf.pages):
            # Layout=True preserves column order (good for both tabular & free-form PDFs)
            text = page.extract_text(
                x_tolerance=3,
                y_tolerance=3,
                layout=True,
                x_density=7.25,
                y_density=13,
            ) or ""

            full_text.append(f"--- PAGE {page_num + 1} ---\n{text}")

    return "\n".join(full_text), page_count


def ocr_pdf(pdf_path: str) -> tuple[str, int]:
    """
    OCR a scanned PDF using pdf2image + pytesseract.
    Returns extracted text and page count.
    """
    poppler_p = get_poppler_path()
    
    # Convert PDF pages to images at 300 DPI for high OCR accuracy
    images = convert_from_path(
        pdf_path,
        dpi=300,
        fmt="PNG",
        thread_count=4,
        poppler_path=poppler_p
    )
    page_count = len(images)
    full_text = []
    
    for i, image in enumerate(images):
        # Convert to grayscale
        gray = image.convert("L")
        # OCR with custom config preserving inter-word spaces for tabular data
        config = "--psm 6 --oem 3 -c preserve_interword_spaces=1"
        try:
            text = pytesseract.image_to_string(gray, lang="eng", config=config)
        except Exception as e:
            # Fallback if tesseract path not configured correctly
            text = f"[OCR Error: Tesseract not configured. Details: {e}]"
            
        full_text.append(f"--- PAGE {i + 1} ---\n{text}")
        
    return "\n".join(full_text), page_count


def extract_text_from_docx(docx_path: str) -> tuple[str, int]:
    """
    Extract text from a .docx file, formatting tables as tab-separated/space-separated rows
    so that they match the TABULAR parser pattern.
    """
    try:
        from docx import Document
        import re
        
        doc = Document(docx_path)
        output_lines = []
        
        body = doc.element.body
        current_team = "BUSINESS OPERATIONS"
        
        paragraphs = {p._element: p for p in doc.paragraphs}
        tables = {t._element: t for t in doc.tables}
        
        MONTHS = ["Jan","Feb","Mar","Apr","May","Jun","Jul","Aug","Sep","Oct","Nov","Dec"]
        
        for child in body:
            if child in paragraphs:
                para = paragraphs[child]
                text = para.text.strip()
                if not text:
                    continue
                
                # Check for team headers, e.g. "Team / Dept: Operations & Events"
                team_match = re.search(r"(?:Team\s*/\s*Dept|Team):\s*([^·\n]+)", text, re.IGNORECASE)
                if team_match:
                    current_team = team_match.group(1).split("Sprint")[0].strip()
                    output_lines.append(f"\nTeam: {current_team}\n")
                else:
                    output_lines.append(text)
                    
            elif child in tables:
                table = tables[child]
                for row in table.rows:
                    cells = [cell.text.strip().replace("\n", " ") for cell in row.cells]
                    if not cells or len(cells) < 4:
                        continue
                    
                    idx_str = cells[0].strip()
                    if not re.match(r"^\d+$", idx_str):
                        continue # Skip headers
                    
                    name = cells[1].strip()
                    task_title = cells[2].strip()
                    priority = cells[3].strip()
                    
                    due_date = ""
                    if len(cells) > 4:
                        raw_due = cells[4].strip()
                        # Normalise to DD-Mon-YY
                        if re.match(r"^\d{1,2}-[A-Za-z]{3}-\d{2}$", raw_due):
                            due_date = raw_due
                        else:
                            dm = re.match(r"(\d{1,2})[/\-](\d{1,2})[/\-](\d{2,4})", raw_due)
                            if dm:
                                try:
                                    d_v, mo_v, yr_v = dm.group(1), int(dm.group(2)), dm.group(3)
                                    if len(d_v) == 1:
                                        d_v = "0" + d_v
                                    due_date = f"{d_v}-{MONTHS[mo_v-1]}-{yr_v[-2:]}"
                                except Exception:
                                    due_date = raw_due
                            else:
                                due_date = raw_due
                    
                    desc = cells[5].strip() if len(cells) > 5 else ""
                    status = cells[6].strip() if len(cells) > 6 else ""
                    
                    # Combine Name + Task Title
                    name_task = f"{name}  {task_title}"
                    
                    # Append Status and Due date as metadata blocks inside Description
                    full_desc = desc
                    if due_date:
                        full_desc += f" (Due: {due_date})"
                    if status:
                        full_desc += f" [Status: {status}]"
                    
                    # Format as standard Tabular row
                    row_line = f"  {idx_str}  {name_task}  {priority}  {full_desc}"
                    output_lines.append(row_line)
                    
        return "\n".join(output_lines), 1
    except Exception as e:
        print(f"[DEBUG] docx tabular extraction failed: {e}. Falling back to docx2txt...")
        try:
            import docx2txt
            text = docx2txt.process(docx_path)
            return text or "", 1
        except Exception as e2:
            raise ValueError(f"Failed to extract text from DOCX: {e2}")


def extract_text_from_doc(doc_path: str) -> tuple[str, int]:
    """
    Extract text from legacy .doc files.
    Tries to convert .doc to .docx using pywin32, then parses using extract_text_from_docx.
    Falls back to binary strings extraction if pywin32/Word is unavailable.
    Returns (text, page_count).
    """
    # 1. Try pywin32 COM automation (highly accurate on Windows with Office)
    try:
        import win32com.client
        import pythoncom
        pythoncom.CoInitialize()
        
        word = win32com.client.Dispatch("Word.Application")
        word.Visible = False
        
        abs_path = os.path.abspath(doc_path)
        temp_docx = abs_path + ".temp.docx"
        
        doc = word.Documents.Open(abs_path)
        doc.SaveAs2(temp_docx, FileFormat=16) # 16 = wdFormatXMLDocument (.docx)
        doc.Close()
        word.Quit()
        
        try:
            # Parse using the docx tabular parser
            text, page_count = extract_text_from_docx(temp_docx)
            return text, page_count
        finally:
            if os.path.exists(temp_docx):
                os.remove(temp_docx)
    except Exception as com_err:
        print(f"[DEBUG] pywin32 COM conversion failed: {com_err}. Using binary fallback...")
        
        # 2. Robust binary parsing fallback (independent of Word/Office)
        try:
            with open(doc_path, "rb") as f:
                content = f.read()
            
            import re
            # Extract plain ASCII text sequences (length >= 4)
            ascii_strings = re.findall(b"[\x20-\x7E\r\n\t]{4,}", content)
            
            # Extract UTF-16LE text sequences (length >= 4 characters)
            utf16_matches = re.findall(b"(?:[\x20-\x7E\r\n\t]\x00){4,}", content)
            
            decoded_utf16 = []
            for match in utf16_matches:
                try:
                    decoded_utf16.append(match.decode("utf-16le"))
                except Exception:
                    pass
            
            decoded_ascii = []
            for match in ascii_strings:
                try:
                    decoded_ascii.append(match.decode("ascii"))
                except Exception:
                    pass
            
            if decoded_utf16:
                utf16_text = "\n".join(decoded_utf16)
                utf16_text = re.sub(r'\n+', '\n', utf16_text)
                if len(utf16_text.strip()) > 50:
                    return utf16_text, 1
            
            ascii_text = "\n".join(decoded_ascii)
            ascii_text = re.sub(r'\n+', '\n', ascii_text)
            return ascii_text, 1
        except Exception as fallback_err:
            raise ValueError(f"Failed to extract text from DOC: {fallback_err}")


